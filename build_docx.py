# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

AR = "Tahoma"

doc = Document()

# default RTL for the whole document (set bidi on styles)
style = doc.styles['Normal']
style.font.name = AR
style.font.size = Pt(12)
rpr = style.element.get_or_add_rPr()
rpr.set(qn('w:bidi'), '1')
style.element.get_or_add_rPr().set(qn('w:bidi'), '1')


def rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p.paragraph_format.alignment = align
    p._p.get_or_add_pPr().set(qn('w:bidi'), '1')
    return p


def H(text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    r.font.name = AR
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    rtl(p)
    return p


def P(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = AR
    r.bold = bold
    rtl(p)
    return p


def BUL(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = AR
    rtl(p)
    p.paragraph_format.left_indent = Pt(20 + level * 18)
    return p


# ---------------- Content ----------------
P("ارسم مخططاً معمارياً واضحاً (Box-and-Arrow Diagram) يوضّح تبادل الصلاحيات والمعلومات والموافقات بين الخادم وهواتف (المدير، المفوض، الموظف المُتتبَّع) بناءً على التوثيق أدناه.", bold=True)

H("المخطط المعماري الصحيح — مواصفات الرسم (Drawing Spec)", 1)
P("ارسم الخادم في المركز كـ«عقدة مركزية»، والهواتف الثلاثة حوله. لا ترسم أي سهم مباشر بين هاتفين؛ كل شيء يمر بالخادم.", bold=True)
P("الصناديق (Boxes):", bold=True)
BUL("هاتف الموظف المُتتبَّع (Tracked Employee).")
BUL("خادم FastAPI + قاعدة البيانات (المنفذ 8001) — في المركز.")
BUL("هاتف المدير / الموارد البشرية (Manager).")
BUL("هاتف المفوض (Viewer).")
P("الأسهم الموجّهة وما تحمله (Directed Edges):", bold=True)
BUL("الموظف ──► الخادم: POST /gps/point (الإحداثيات lat/lng فقط). لا سهم من الموظف للمدير مباشرة.")
BUL("المدير ──► الخادم: POST /gps/start و /gps/stop و /gps/set-viewer (أوامر تحكّم فقط).")
BUL("المفوض ──► الخادم: GET /gps/view و /gps/history و تصدير PDF (مشاهدة فقط، لا تحكّم).")
BUL("الخادم ──► المدير/المفوض: GET /gps/view + /history (يسحبان بيانات المسير، لا تُرسَل تلقائياً).")
BUL("الخادم ──► الموظف + المدير + المفوض: WebSocket /ws/notifications (تنبيهات الخروج/العودة فقط — لا إحداثيات).")
P("تبادل الصلاحيات: يُمثَّل بالتوكن JWT عند الدخول (لا كرسالة بين الهواتف). تبادل الموافقات: يبدأه المدير (start) وينهيه (stop) عبر الخادم.", bold=True)

H("تحذيرات لتجنّب الأخطاء الشائعة عند الرسم", 2)
BUL("خطأ: سهم موقع من الموظف مباشرة إلى المدير. الصحيح: يمر بالخادم (الموظف→خادم→مدير).")
BUL("خطأ: جعل WebSocket ينقل الإحداثيات. الصحيح: WebSocket ينقل التنبيهات فقط.")
BUL("خطأ: إعطاء المفوض زر بدء/إيقاف. الصحيح: المفوض يراقب ويطبع فقط.")
BUL("خطأ: المدير يرسل الأوامر مباشرة للموظف. الصحيح: عبر الخادم فقط.")
BUL("خطأ: الموظف يرى خريطة الآخرين. الصحيح: الموظف يرى شريط حالته فقط.")
BUL("خطأ: توكن الصلاحية يُرسَل بين الهواتف. الصحيح: يُرسَل من كل هاتف للخادم في ترويسة Authorization.")

H("نظام تتبّع المواقع — التوثيق النصي الكامل", 1)

H("٠) المكوّنات الثلاثة", 2)
BUL("الخادم (Backend): FastAPI + SQLAlchemy(async) + قاعدة SQLite، يُشغَّل بـ uvicorn على المنفذ 8001، والمصادقة بنظام JWT Bearer.")
BUL("تطبيق الهاتف (Flutter): عربي RTL، يعمل على ثلاثة أدوار: المدير/الموارد البشرية، الموظف المُتتبَّع، والموظف المفوض (Viewer).")
BUL("قناة التنبيهات: WebSocket على /ws/notifications تنقل التنبيهات فوراً.")
BUL("كل التواصل بين الأطراف يتم عبر REST (HTTP/HTTPS) للبيانات، وWebSocket للتنبيهات اللحظية.")

H("١) الأدوار والصلاحيات (RBAC)", 2)
P("المدير العام / موارد بشرية (gm/hr): يبدأ/يوقف التتبّع، يعيّن المفوض، يرى كل الجلسات والمسير، يرى سجل الخروقات، يستقبل تنبيهات الخروج. لا يُرسل موقعاً.")
P("موظف مُتتبَّع (employee): يرسل نقاط موقعه، يرى حالته فقط، يستقبل تنبيه «خارج المنطقة». لا يبدأ تتبّعاً ولا يرى خريطة الآخرين.")
P("مفوض (Viewer): يرى مسير أي موظف، يطبع PDF، يستقبل تنبيهات الخروج والملخّص. لا يبدأ/يوقف ولا يعيّن مفوضاً.")
P("كيف تُتحقَّق الصلاحية: عند كل طلب يُرسل التطبيق التوكن Authorization: Bearer <token>. الخادم يفك التوكن ويعرف الهوية والدور، ثم يعتمد أو يرفض الطلب بحسب الصلاحية المطلوبة.")

H("٢) تبادل الصلاحيات (كيف تُمنح وتُتحقَّق)", 2)
BUL("منح الصلاحية: عند تسجيل الدخول عبر POST /auth/login {employee_number, password} يُرجع الخادم access_token (JWT يحمل الهوية + الدور). هذا التوكن هو رخصة العمل لكل الطلبات.")
BUL("التحقّق المستمر: كل طلب GPS يمر عبر وسيط يقرأ التوكن ويقرر السماح.")
BUL("POST /gps/start, stop, set-viewer, simulate-point → يسمح فقط لـ gm/hr.")
BUL("POST /gps/point → يسمح لأي موظف، لكن employee_id النقطة يجب = صاحب التوكن (لا يستطيع موظف إرسال موقع موظف آخر).")
BUL("GET /gps/view, history, breaches, employees, viewer → يسمح لـ gm/hr والمفوض فقط.")
BUL("GET /gps/my-active → يسمح لأي موظف (يرى حالته هو فقط).")
BUL("صلاحية المفوض: لا دور ثابت، بل تُحدَّد وقت التشغيل عبر AppSetting.tracking_viewer_id؛ أي موظف يُعيَّن مفوضاً يكتسب صلاحية المشاهدة/الطباعة فوراً.")

H("٣) تبادل المعلومات بين الأطراف", 2)
P("أ) من المدير → الخادم:", bold=True)
BUL("POST /gps/set-viewer {employee_id} → يعيّن المفوض (يُحفظ في app_settings).")
BUL("POST /gps/start {target_employee_id, mode, interval} → ينشئ FieldTrackingSession (status=active).")
BUL("POST /gps/stop {session_id} → يُنهي الجلسة (status=completed).")
BUL("POST /gps/simulate-point {session_id, latitude, longitude} → للاختبار فقط.")
P("ب) من الموظف → الخادم:", bold=True)
BUL("GET /gps/my-active → يسحب إن كان مُتتبَّعاً (session_id + mode + interval + is_outside + outside_distance_m).")
BUL("POST /gps/point {session_id, latitude, longitude, accuracy, speed, battery_level} → يُخزَّن في field_tracking_points ثم يُشغَّل كشف geofence.")
P("ج) من الخادم → المدير/المفوض:", bold=True)
BUL("GET /gps/view → كل الجلسات النشطة + آخر نقطة + geofence + is_outside + outside_started_at + outside_distance_m.")
BUL("GET /gps/history?employee_id= → نقاط المسير (لرسم الخط على الخريطة).")
BUL("GET /gps/employees → دليل الموظفين مع حقول منطقتهم.")
BUL("GET /gps/breaches?employee_id= → سجل الخروقات (للتقرير/PDF).")
BUL("GET /gps/viewer → من هو المفوض الحالي.")
P("د) من الخادم → الموظف:", bold=True)
BUL("GET /gps/my-active → حالته (داخل/خارج + المدة + المسافة) لعرض الشريط الحي.")

H("٤) الموافقات داخل التطبيق (Approvals)", 2)
BUL("موافقة البدء (Authorization to track): عندما يضغط المدير «بدء» فهو يُصدر تفويضاً رسمياً بتتبّع الموظف لفترة الجلسة؛ الخادم يُنشئ الجلسة ويُفعّل إرسال الموقع من هاتف الموظف. لا يستطيع الموظف رفضها داخل التطبيق — التتبّع يبدأ تلقائياً بمجرد أن يسحب هاتفه my-active ويجد جلسة نشطة.")
BUL("موافقة تعيين المفوض: المدير يعيّن موظفاً بصلاحية المراقبة/الطباعة؛ توافق إداري يُحفظ فوراً.")
BUL("موافقة الإيقاف: المدير يوقف الجلسة → ينتهي التفويض → هاتف الموظف يتوقف عن الإرسال.")
BUL("معالجة الخروج (Breach): عند خروج الموظف من منطقته يولّد النظام «خرق» (GeofenceBreach) بعد عودته، ويرسل ملخّصاً للمديرين والمفوض كسجل موثَّق آلياً.")

H("٥) تدفّق البيانات خطوة بخطوة (Sequence)", 2)
BUL("1) المدير: POST /gps/start → الخادم ينشئ FieldTrackingSession (active).")
BUL("2) المدير: يسحب GET /gps/view كل 5 ثوانٍ → تظهر الخريطة مع «بانتظار أول نقطة».")
BUL("3) الموظف: LocationTrackerService يسحب GET /gps/my-active كل 15 ثانية → يرى active=true + session_id + mode.")
BUL("4) الموظف: يرسل POST /gps/point (lat/lng) بالوتيرة (مسافة أو زمن) التي حدّدها المدير.")
BUL("5) الخادم: يخزّن النقطة، ويحسب المسافة إلى مركز منطقة العمل (Haversine).")
BUL("6) إن كان خارجاً: الخادم ينشئ إشعاراً للموظف («أنت خارج منطقة عملك») + للمدير + للمفوض («خرج خارج منطقة عمله») ويدفعه فوراً عبر WebSocket، ويبدأ عدّ الزمن والمسافة (outside_distance_m).")
BUL("7) المدير/المفوض: GET /gps/view يُظهر is_outside=true + outside_started_at + outside_distance_m → شارة حية بعدّاد.")
BUL("8) الموظف: GET /gps/my-active يُظهر is_outside → شريط «أنت خارج منطقة عملك» بعدّاد.")
BUL("9) عند العودة: الخادم ينشئ سجل GeofenceBreach (المدة + المسافة) + إشعار ملخّص للمديرين والمفوض.")
BUL("10) المدير: POST /gps/stop → الجلسة completed → الموظف يتوقف عن الإرسال.")
BUL("11) أي وقت: المدير/المفوض يفتح GET /gps/breaches → تقرير + تصدير PDF.")

H("٦) التنبيهات الفورية (WebSocket)", 2)
BUL("المسار: ws://<server>/ws/notifications?token=<JWT> (يصبح wss:// تلقائياً مع https).")
BUL("connection_manager يربط employee_id بقائمة اتصالاته النشطة.")
BUL("عند الخروج/العودة: notify_employee / notify_managers تُنشئ صفّاً في جدول notifications وتدفع JSON فوراً عبر WS للهواتف المعنية.")
BUL("حمولة التنبيه: {notification_id, title, message, severity, created_at}.")
BUL("التطبيق (notification_ws) يعرض الإشعار ويحدّث الشارة دون سحب البيانات.")

H("٧) الشبكة والربط", 2)
BUL("داخل المؤسسة (نفس الواي فاي): عنوان الخادم بـ⚙ = http://192.168.1.18:8001.")
BUL("بُعد 100 كم / شبكة مختلفة: الخادم يُكشَف للإنترنت عبر نفق cloudflared/ngrok (تجريباً) أو استضافة VPS/سحابة (إنتاجياً) برابط https عام.")
BUL("التطبيق يحوّل http→ws و https→wss تلقائياً (لا حاجة لتعديل كود).")
BUL("شرط أساسي: فتح المنفذ 8001 بالجدار الناري (دخول)، وعمل الخادم دائماً.")

H("٩) نظرة عامة على التطبيق بأكمله (الوحدات)", 1)
P("النظام منصة متكاملة لإدارة مياه يارموك/عجلون، ومكوّناته:", bold=True)
BUL("خادم خلفي FastAPI (المنفذ 8001) + قاعدة SQLite.")
BUL("تطبيق هاتف Flutter (عربي RTL) للموظفين والمديرين والمفوض.")
BUL("لوحة تحكم ويب (Web Dashboard) للإدارة.")
BUL("قناة تنبيهات WebSocket فورية.")
P("وحدات الخادم (Backend Routers):", bold=True)
BUL("auth: تسجيل الدخول + JWT + الأدوار والصلاحيات.")
BUL("employees: إدارة الموظفين + تعريف منطقة العمل (geofence) + الصلاحيات.")
BUL("organization: الهيكل التنظيمي والأقسام.")
BUL("attendance: الحضور والانصراف.")
BUL("leave_requests: طلبات الإجازات مع دورة موافقة.")
BUL("overtime_work: العمل الإضافي مع دورة موافقة.")
BUL("violations: المخالفات.")
BUL("maintenance: طلبات/أعمال الصيانة.")
BUL("customer_service: خدمة العملاء.")
BUL("water_distribution: توزيع المياه.")
BUL("work_scopes: نطاقات العمل.")
BUL("gps: التتبّع + المسير + كشف الخروق (موثّق أعلاه).")
BUL("notifications: الإشعارات (REST + WebSocket).")
BUL("reports: التقارير.  audit: تدقيق العمليات.")
P("شاشات تطبيق الهاتف:", bold=True)
BUL("الدخول، اللوحة الرئيسية، الملف الشخصي.")
BUL("التتبّع GPS، التتبّع الحي، سجل الخروق.")
BUL("الموظفون، الهيكل التنظيمي.")
BUL("الحضور، الإجازات، العمل الإضافي، المخالفات، الصيانة، التقارير، صندوق التقارير، الإشعارات.")

H("١٠) المعمارية الشاملة للنظام (Overall Architecture)", 2)
P("ارسم مخططاً طبقات (Layered Architecture Diagram) يوضّح النظام ككل:", bold=True)
BUL("طبقة العرض (Presentation): تطبيق الهاتف (Flutter) + لوحة الويب — عربي RTL.")
BUL("طبقة الواجهة (API): FastAPI REST + WebSocket على المنفذ 8001.")
BUL("طبقة الأعمال (Services): صلاحيات، تنبيهات، تتبّع GPS.")
BUL("طبقة البيانات (Data): SQLAlchemy(async) + قاعدة SQLite.")
BUL("التدفّق العام: العميل (هاتف/ويب) → HTTPS/WS → FastAPI → Services → DB. كل طلب يحمل توكن JWT؛ الصلاحيات حسب الدور.")
BUL("مخطط الصناديق المقترح: [تطبيق الهاتف] و [لوحة الويب] ──► [FastAPI API + WebSocket] ──► [Services] ──► [SQLite]؛ وقناة WebSocket تعود من الخادم إلى العملاء للتنبيهات الفورية.")
BUL("وحدة GPS (الموثّقة في الأقسام ١–٨) هي أحد وحدات الطبقة نفسها، وتتبع نفس نموذج الصلاحيات والمصادقة.")

H("١١) خلاصة", 2)
P("النظام يتبادل الصلاحيات عبر التوكن والدور، المعلومات عبر REST (نقاط/جلسات/مسير/خروق)، والموافقات عبر تفويض المدير بالبدء/التعيين/الإيقاف وتوثيق الخروق آلياً. كل ذلك داخل التطبيق بصلاحيات صارمة: المدير يتحكّم، الموظف يُرسل موقعه فقط، والمفوض يراقب ويطبع.")

doc.save(r"D:\yarmouk_water_management_pro\gps_documentation.docx")
print("SAVED gps_documentation.docx paragraphs:", len(doc.paragraphs))
